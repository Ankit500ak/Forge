from datetime import datetime
from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.2

    heading1 = doc.styles["Heading 1"]
    heading1.font.name = "Calibri"
    heading1.font.size = Pt(16)
    heading1.font.bold = True

    heading2 = doc.styles["Heading 2"]
    heading2.font.name = "Calibri"
    heading2.font.size = Pt(13)
    heading2.font.bold = True

    heading3 = doc.styles["Heading 3"]
    heading3.font.name = "Calibri"
    heading3.font.size = Pt(12)
    heading3.font.bold = True


def add_cover_page(doc: Document, title: str, subtitle: str) -> None:
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_title = p_title.add_run(title)
    r_title.bold = True
    r_title.font.size = Pt(24)

    p_subtitle = doc.add_paragraph()
    p_subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_subtitle = p_subtitle.add_run(subtitle)
    r_subtitle.font.size = Pt(14)

    doc.add_paragraph("")
    p_project = doc.add_paragraph()
    p_project.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_project.add_run("Project: FORGE Fitness Platform").bold = True

    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_date.add_run(f"Prepared on: {datetime.now().strftime('%d %B %Y')}")

    doc.add_page_break()


def sanitize_line(line: str) -> str:
    cleaned = line.strip()
    cleaned = re.sub(r"\*\*(.*?)\*\*", r"\1", cleaned)
    cleaned = re.sub(r"`([^`]*)`", r"\1", cleaned)
    return cleaned


def add_markdown_body(doc: Document, text: str) -> None:
    lines = text.splitlines()
    for raw_line in lines:
        line = raw_line.rstrip()
        if not line.strip():
            doc.add_paragraph("")
            continue

        heading_match = re.match(r"^(#{1,6})\s+(.*)$", line)
        if heading_match:
            level = min(len(heading_match.group(1)), 3)
            text_value = sanitize_line(heading_match.group(2))
            doc.add_heading(text_value, level=level)
            continue

        bullet_match = re.match(r"^[-*]\s+(.*)$", line)
        if bullet_match:
            p = doc.add_paragraph(sanitize_line(bullet_match.group(1)), style="List Bullet")
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.15
            continue

        numbered_match = re.match(r"^\d+[\.)]\s+(.*)$", line)
        if numbered_match:
            p = doc.add_paragraph(sanitize_line(numbered_match.group(1)), style="List Number")
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.15
            continue

        p = doc.add_paragraph(sanitize_line(line))
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.2


def convert(md_path: Path, docx_path: Path, subtitle: str) -> None:
    content = md_path.read_text(encoding="utf-8")
    doc = Document()
    configure_document(doc)

    first_line = next((ln for ln in content.splitlines() if ln.strip()), md_path.stem)
    title = sanitize_line(first_line.replace("#", "").strip())

    add_cover_page(doc, title=title, subtitle=subtitle)
    add_markdown_body(doc, content)
    doc.save(docx_path)


def create_combined_dossier(root: Path, output_path: Path) -> None:
    doc = Document()
    configure_document(doc)
    add_cover_page(
        doc,
        title="FORGE Fitness Platform",
        subtitle="Professional Project Dossier",
    )

    sources = [
        root / "PROJECT_INFORMATION.md",
        root / "PROFESSIONAL_SYNOPSIS.md",
    ]

    for index, source in enumerate(sources):
        content = source.read_text(encoding="utf-8")
        section_title = source.stem.replace("_", " ").title()
        doc.add_heading(section_title, level=1)
        add_markdown_body(doc, content)
        if index < len(sources) - 1:
            doc.add_page_break()

    doc.save(output_path)


def main() -> None:
    root = Path(__file__).resolve().parent.parent
    pairs = [
        (
            root / "PROJECT_INFORMATION.md",
            root / "PROJECT_INFORMATION.docx",
            "Comprehensive Information File",
        ),
        (
            root / "PROFESSIONAL_SYNOPSIS.md",
            root / "PROFESSIONAL_SYNOPSIS.docx",
            "Formal Synopsis Document",
        ),
    ]

    for md_path, docx_path, subtitle in pairs:
        if not md_path.exists():
            raise FileNotFoundError(f"Missing source file: {md_path}")
        convert(md_path, docx_path, subtitle)
        print(f"Created: {docx_path}")

    dossier_path = root / "FORGE_PROFESSIONAL_DOSSIER.docx"
    create_combined_dossier(root, dossier_path)
    print(f"Created: {dossier_path}")


if __name__ == "__main__":
    main()
